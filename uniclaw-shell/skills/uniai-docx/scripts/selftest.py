"""自检：用每个 preset 各生成一份样张，跑规则检查，验证关键 XML 不变量。

    python scripts/selftest.py            # 临时目录，跑完即删
    python scripts/selftest.py --keep out # 把样张留在 out/ 里，可以拿去渲染

改动 docx_helpers.py / check_docx.py / SKILL.md 之后跑这个——它是“文档说的”和
“代码做的”开始分叉时的报警器。不依赖 LibreOffice。
"""

from __future__ import annotations

import argparse
import sys
import tempfile
from os.path import abspath, dirname, exists, join

sys.path.insert(0, dirname(abspath(__file__)))

# Windows 控制台默认 GBK，编不出的字符会让 print 抛 UnicodeEncodeError 中断脚本。
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(errors="replace")

from docx import Document                                    # noqa: E402
from docx.oxml.ns import qn                                  # noqa: E402

import check_docx as ck                                      # noqa: E402
from docx_helpers import (Cm, PRESETS, Pt, THEMES, add_cover, add_figure,   # noqa: E402
                          add_landscape_section, add_portrait_section, add_toc,
                          new_document, start_body_section, style_table)

ASSET = join(dirname(dirname(abspath(__file__))), "assets", "uniai-docx.png")
W = ck.W
FAILURES: list[str] = []


def check(cond, msg):
    print(("  ok   " if cond else "  FAIL ") + msg)
    if not cond:
        FAILURES.append(msg)


def build(path, theme, preset, with_figure=True):
    doc = new_document(theme=theme, preset=preset)
    add_cover(doc, "2026 年上半年经营分析报告", "市场与运营中心",
              ["编制单位：XX 科技有限公司", "2026 年 8 月"])
    start_body_section(doc, page_number=True, header_text="经营分析报告")
    add_toc(doc)

    doc.add_heading("一、总体经营情况", level=1)
    doc.add_paragraph("上半年公司整体营收保持稳健增长，各业务线表现分化，具体情况如下所述。")
    doc.add_heading("1.1 收入结构", level=2)
    doc.add_paragraph("从收入构成看，核心业务仍是主要来源，新业务占比稳步提升。")

    data = [["业务线", "收入（万元）", "同比增速", "占比"],
            ["核心业务", "12,480", "8.3%", "54.2%"],
            ["新兴业务", "5,320", "41.7%", "23.1%"],
            ["海外业务", "2,910", "19.4%", "12.6%"],
            ["服务与运维", "1,540", "5.1%", "6.7%"],
            ["硬件销售", "520", "-3.2%", "2.3%"],
            ["其他", "250", "1.1%", "1.1%"],
            ["合计", "23,020", "13.6%", "100.0%"],
            ["去年同期", "20,260", "13.6%", "100.0%"]]
    t = doc.add_table(rows=len(data), cols=4)
    for r, row in enumerate(data):
        for c, v in enumerate(row):
            t.cell(r, c).text = v
    style_table(t, doc=doc, align_right_cols=(1, 2, 3))

    if with_figure and exists(ASSET):
        add_figure(doc, ASSET, "核心业务贡献过半收入", width_in=2.0, source="公司财报")

    doc.add_heading("二、风险与建议", level=1)
    doc.add_paragraph("硬件销售出现小幅下滑，需要在下半年调整渠道策略。")
    doc.save(path)
    return path


def rules_pass(path):
    """跑 check_docx 的全部检查，返回 ERROR 列表。"""
    doc = Document(path)
    rep = ck.Report()
    smap = ck.style_font_map(doc)
    body_w = ck.check_page(doc, rep, False)
    ck.check_fonts(doc, rep, smap)
    sizes = ck.check_runs(doc, rep, smap)
    heads = ck.check_structure(doc, rep, body_w, sizes)
    ck.check_characters(doc, rep)
    ck.check_tables(doc, rep)
    ck.check_fields(doc, rep, heads)
    return [i for i in rep.items if i["level"] == "ERROR"], \
           [i for i in rep.items if i["level"] == "WARN"]


def invariants(path, preset):
    doc = Document(path)
    sec = doc.sections[0]
    expect_a4 = preset != "letter"
    if expect_a4:
        check(round(sec.page_width.cm, 1) == 21.0 and round(sec.page_height.cm, 1) == 29.7,
              "页面是 A4")

    styles_xml = doc.styles.element.xml
    check("w:eastAsia" in styles_xml, "样式层写入了 eastAsia 中文字体")
    for sid in ("Normal", "Heading1", "Heading2", "Heading3"):
        node = ck.style_font_map(doc).get(sid, {})
        check(not (node.get("theme") and node.get("eastAsia")),
              f"{sid} 已清掉主题字体属性（否则显式字体不生效）")

    tbl = doc.tables[0]._tbl
    check(tbl.tblPr.find(qn("w:tblBorders")) is not None, "表格有显式边框（非黑网格）")
    check(tbl.tblPr.find(qn("w:tblCellMar")) is not None, "表格设了单元格内边距")
    check(doc.tables[0].rows[0]._tr.find(f"{W}trPr/{W}tblHeader") is not None,
          "长表格设了跨页表头重复")
    check(doc.tables[0].cell(0, 0)._tc.find(f"{W}tcPr/{W}shd") is not None, "表头有底色")

    check(len(doc.sections) >= 2, "封面与正文分属不同节")
    starts = [s._sectPr.find(qn("w:pgNumType")) for s in doc.sections]
    restarts = [s for s in starts[1:] if s is not None and s.get(qn("w:start")) == "1"]
    check(len(restarts) == 1, "正文页码从 1 重新开始，且只重编一次")
    check(any("PAGE" in " ".join(p._p.xml for p in s.footer.paragraphs)
              for s in doc.sections[1:]), "正文页脚有页码域")
    check("PAGE" not in " ".join(p._p.xml for p in doc.sections[0].footer.paragraphs),
          "封面页脚没有页码")

    # 换页一律用分节符：不留空段落，预览里也不会出现分页符方块
    check(not any(("w:br" in p._p.xml and 'w:type="page"' in p._p.xml)
                  for p in doc.paragraphs), "换页不靠空段落 + 手工分页符")
    check(not any(not p.text.strip() and "w:pBdr" not in p._p.xml
                  and "w:sectPr" not in p._p.xml and "w:drawing" not in p._p.xml
                  for p in doc.paragraphs), "正文里没有空段落")

    check("TOC" in doc.element.xml, "插入了目录域")


def skill_doc_consistency():
    """SKILL.md 里承诺的东西，代码里得真有。"""
    root = dirname(dirname(abspath(__file__)))
    skill = open(join(root, "SKILL.md"), encoding="utf-8").read()
    import docx_helpers as dh
    for name in ("new_document", "add_cover", "start_body_section", "add_toc",
                 "style_table", "add_figure", "set_font", "apply_theme",
                 "add_landscape_section", "add_portrait_section", "set_columns"):
        if name in skill:
            check(hasattr(dh, name), f"SKILL.md 提到的 {name}() 在 docx_helpers 里存在")
    for f in ("scripts/check_docx.py", "scripts/render_docx.py", "scripts/docx_helpers.py",
              "references/design-taste.md"):
        if f in skill:
            check(exists(join(root, f)), f"SKILL.md 引用的 {f} 存在")
    for theme in THEMES:
        check(theme in skill, f"主题 {theme} 在 SKILL.md 里有说明")
    for preset in PRESETS:
        check(preset in skill, f"预设 {preset} 在 SKILL.md 里有说明")


CROSS_PLATFORM_FONTS = {"宋体", "黑体", "楷体", "仿宋", "仿宋_GB2312",
                        "Times New Roman", "Arial"}


def font_portability():
    """预设里的字体必须 Windows / macOS 都有，否则同一份文档两台机器长相不同。"""
    for name, p in PRESETS.items():
        for key in ("latin", "cjk", "h_latin", "h_cjk"):
            font = p.get(key)
            if not font:
                continue
            ok = font in CROSS_PLATFORM_FONTS or name == "gongwen"   # 公文按国标，Windows 为准
            check(ok, f"预设 {name} 的 {key}='{font}' 是跨平台字体")


def landscape_sample(path):
    """横竖混排样张：竖版正文 + 横版宽表节 + 切回竖版。"""
    doc = new_document(theme="teal", preset="business")
    doc.add_heading("一、竖版正文", level=1)
    doc.add_paragraph("这一节是连续文字，应当保持竖版。" * 4)

    add_landscape_section(doc)
    doc.add_heading("二、宽表（横版节）", level=1)
    t = doc.add_table(rows=3, cols=9)
    for r in range(3):
        for c in range(9):
            t.cell(r, c).text = f"指标{c+1}" if r == 0 else f"{(r * 9 + c) * 1370:,}"
    style_table(t, doc=doc, align_right_cols=tuple(range(1, 9)))

    add_portrait_section(doc)
    doc.add_heading("三、回到竖版", level=1)
    doc.add_paragraph("后续正文恢复竖版。")
    doc.save(path)

    d = Document(path)
    orients = [("横" if s.page_width.cm > s.page_height.cm else "竖") for s in d.sections]
    check(orients[0] == "竖" and "横" in orients and orients[-1] == "竖",
          f"横竖混排正确（{'/'.join(orients)}）：正文竖版、宽表横版、之后切回")
    land = next(s for s in d.sections if s.page_width.cm > s.page_height.cm)
    check(round(land.page_width.cm, 1) == 29.7 and round(land.page_height.cm, 1) == 21.0,
          "横版节是 A4 横版 29.7×21")
    errors, _ = rules_pass(path)
    check(not errors, "横竖混排样张 check_docx 无 ERROR" +
          ("" if not errors else "：" + "；".join(e["code"] for e in errors)))


def page_estimate_calibration(tmp):
    """页数估算要可信：高估会把智能体推去做没必要的缩字号。

    基准来自渲染实测：一份 2359 字的周报，横版 A4 / 边距 1.2cm / 列宽按内容分配，
    8pt 实测占约 0.68 页。这里用同量级的合成内容复现该量级，并断言随字号单调。
    """
    import re as _re
    from docx_helpers import _set_cell_margins

    def build(pt, path):
        doc = new_document(theme="navy", preset="modern", orientation="landscape")
        s = doc.sections[0]
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.2)
        doc.styles["Normal"].paragraph_format.line_spacing = 1.1
        doc.styles["Normal"].paragraph_format.space_after = Pt(0)
        cell_text = "①" + "任务进展说明内容" * 12 + "；②" + "后续计划与风险提示" * 10 + "。"
        t = doc.add_table(rows=9, cols=5)
        for c, v in enumerate(["任务", "工作进展", "需协调问题", "责任人A", "责任人B"]):
            t.cell(0, c).text = v
        for r in range(1, 9):
            t.cell(r, 0).text = f"板块{r}"
            t.cell(r, 1).text = cell_text
            t.cell(r, 2).text = "1、需要对方确认；2、等待资源到位"
            t.cell(r, 3).text = "张三 13800000000"
            t.cell(r, 4).text = "李四 13900000000"
        style_table(t, doc=doc, font_size=pt, zebra=True)
        _set_cell_margins(t, top=40, left=80, bottom=40, right=80)
        doc.save(path)
        return path

    ratios = {}
    for pt in (8, 10.5):
        d = Document(build(pt, join(tmp, f"calib_{pt}.docx")))
        rep = ck.Report()
        pages = ck.check_layout(d, rep, None)
        msg = next(i["msg"] for i in rep.items if i["code"] == "PAGE_ESTIMATE")
        ratios[pt] = (pages, float(_re.search(r"([\d.]+) 个版心", msg).group(1)))

    check(ratios[8][0] == 1,
          f"8pt 的密排表格估算为 1 页（实得 {ratios[8][0]} 页 / {ratios[8][1]} 版心）")
    check(ratios[10.5][1] > ratios[8][1],
          f"字号变大时估算高度单调增加（8pt={ratios[8][1]} < 10.5pt={ratios[10.5][1]}）")
    check(0.4 < ratios[8][1] < 1.0,
          f"8pt 估算落在合理量级（{ratios[8][1]} 版心，实测基准 0.68）")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--keep", default=None, help="把样张留在这个目录")
    args = ap.parse_args()

    tmp = args.keep or tempfile.mkdtemp(prefix="docx_selftest_")
    import os
    os.makedirs(tmp, exist_ok=True)

    print("SKILL.md ↔ 代码一致性")
    skill_doc_consistency()
    print("\n字体跨平台性")
    font_portability()
    print("\n横竖混排")
    landscape_sample(join(tmp, "sample_landscape.docx"))
    print("\n页数估算校准")
    page_estimate_calibration(tmp)

    for preset in PRESETS:
        theme = list(THEMES)[list(PRESETS).index(preset) % len(THEMES)]
        path = join(tmp, f"sample_{preset}.docx")
        print(f"\npreset={preset} theme={theme} → {path}")
        build(path, theme, preset)
        errors, warns = rules_pass(path)
        check(not errors, "check_docx 无 ERROR" +
              ("" if not errors else "：" + "；".join(e["code"] for e in errors)))
        if warns:
            print("       (warn: " + "、".join(w["code"] for w in warns) + ")")
        invariants(path, preset)

    print()
    if FAILURES:
        print(f"× {len(FAILURES)} 项未通过：")
        for f in FAILURES:
            print("   - " + f)
        sys.exit(1)
    print(f"√ 全部通过。样张在 {tmp}"
          + ("" if args.keep else "（临时目录，可用 --keep 保留）"))


if __name__ == "__main__":
    main()
