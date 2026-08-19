"""python-docx 排版助手：把 SKILL.md 的硬规则变成可调用的函数。

用法（在 workspace 里）：
    import sys; sys.path.insert(0, '<skill_dir>/scripts')
    from docx_helpers import *

    doc = new_document(theme='navy', preset='business')      # A4 + 样式 + 字体，一步到位
    add_cover(doc, '2026 年度经营分析报告', '市场部', ['编制单位：XX 公司', '2026 年 8 月'])
    start_body_section(doc, page_number=True)                 # 正文另起节，页码从 1 开始
    add_toc(doc)
    doc.add_heading('一、总体情况', level=1)
    p = doc.add_paragraph('正文……')

为什么要用这些函数而不是裸写 python-docx：
  * python-docx 默认页面是 US Letter，不是 A4；
  * 它设置字体时保留 w:asciiTheme/w:eastAsiaTheme，主题属性会覆盖你设的字体；
  * 新建 run 的 rPr 是 None，直接 `run._element.rPr.rFonts` 会抛 AttributeError；
  * 默认模板里唯一带边框的表样式是 Table Grid（全黑网格）。
上面四点每一点都会让文档“看起来不对”，这里全部封装掉了。
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

import re

CJK_RE = re.compile(r"[㐀-鿿豈-﫿　-〿＀-￯]")

__all__ = [
    "THEMES", "PRESETS", "new_document", "setup_page", "apply_theme",
    "set_font", "set_first_line_indent_chars", "add_cover", "start_body_section",
    "add_page_number", "add_toc", "add_field", "add_rule", "add_figure",
    "style_table", "set_column_alignment", "shade_cell", "hex_rgb",
    "add_landscape_section", "add_portrait_section", "set_columns",
    "body_width_cm", "fit_table_to_page", "page_break_section",
    "auto_column_widths",
    # 常用单位与色值一并导出，省得再从 docx.shared 单独 import
    "Cm", "Pt", "Inches", "RGBColor", "INK", "MUTED", "LINE", "ZEBRA",
]

# --- 主题色板（与 SKILL.md 的五套主题一一对应） -----------------------------
THEMES = {
    "navy":   {"accent": "1F4E79", "tint": "EAF1F8", "name": "藏青·稳重"},
    "green":  {"accent": "1E5C40", "tint": "E8F2EC", "name": "墨绿·信实"},
    "teal":   {"accent": "0F766E", "tint": "E5F2F0", "name": "青碧·现代"},
    "ochre":  {"accent": "9A3412", "tint": "FAEEE7", "name": "赭石·人文"},
    "claret": {"accent": "8A1C2B", "tint": "F9ECEE", "name": "酒红·庄重"},
}
INK = "1F2937"        # 正文墨色，别用纯黑
MUTED = "595959"      # 弱化灰：图注、页眉页脚、来源
LINE = "D9D9D9"       # 表格内框线
ZEBRA = "F7F8FA"      # 斑马纹

# --- 版式预设 ---------------------------------------------------------------
# body/h1/h2/h3 单位 pt；line 为倍数行距，line_exact 为固定值行距（pt，优先）。
#
# 字体只用 Windows 与 macOS 都有的：黑体 / 宋体 / 仿宋 / 楷体 + Times New Roman / Arial。
# 微软雅黑、Segoe UI 是 Windows 独有的，macOS 上会静默回退成别的字体，
# 同一份文档在两台机器上长相不同——这是实测踩到的坑（macOS 预览里标题变成了衬线体）。
# 明确只在 Windows 交付时，可以传 preset 字典覆盖 h_cjk='微软雅黑'、cjk='微软雅黑'。
PRESETS = {
    # 商务报告：标题黑体、正文宋体，小四 12pt / 1.5 倍行距
    "business": {
        "latin": "Times New Roman", "cjk": "宋体",
        "h_latin": "Arial", "h_cjk": "黑体",
        "body": 12, "h1": 16, "h2": 14, "h3": 12,
        "line": 1.5, "space_after": 8, "indent_chars": 2,
    },
    # 现代商务：全文无衬线，五号 10.5pt，信息密度更高（Windows 上可换 微软雅黑 更佳）
    "modern": {
        "latin": "Arial", "cjk": "黑体",
        "h_latin": "Arial", "h_cjk": "黑体",
        "body": 10.5, "h1": 15, "h2": 13, "h3": 11,
        "line": 1.45, "space_after": 8, "indent_chars": 0,
    },
    # 公文：GB/T 9704，标题黑体、正文仿宋三号，每页 22 行
    # 仿宋_GB2312 是国标要求的字体，Windows 独有；macOS 上会回退，公文本就以 Windows 交付为准
    "gongwen": {
        "latin": "Times New Roman", "cjk": "仿宋_GB2312",
        "h_latin": "Times New Roman", "h_cjk": "黑体",
        "body": 16, "h1": 16, "h2": 16, "h3": 16,
        "line": None, "line_exact": 29, "space_after": 0, "indent_chars": 2,
        "margins": (3.7, 2.6, 3.5, 2.8), "grid_lines": 22,
    },
    # 学术/说明书：正文宋体五号，紧凑但留白足
    "academic": {
        "latin": "Times New Roman", "cjk": "宋体",
        "h_latin": "Times New Roman", "h_cjk": "黑体",
        "body": 10.5, "h1": 15, "h2": 13, "h3": 11,
        "line": 1.5, "space_after": 6, "indent_chars": 2,
    },
}


# --- 底层工具 ---------------------------------------------------------------
def _el(tag: str, **attrs) -> OxmlElement:
    """建一个 w: 元素，属性名自动补 w: 前缀。"""
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def hex_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _rpr(obj):
    """取 run / style 的 rPr，没有就建——避开 rPr 为 None 的崩溃。"""
    el = getattr(obj, "_element", None)
    if el is None:
        el = getattr(obj, "element", None)
    if el is None:
        raise TypeError(f"不支持的对象：{type(obj)}")
    return el.get_or_add_rPr()


def set_font(obj, latin=None, cjk=None, size=None, bold=None,
             italic=None, color=None):
    """给 run 或 style 设字体。**必须用这个函数，不要裸写 rFonts。**

    关键：清掉 w:asciiTheme / w:eastAsiaTheme 等主题属性。OOXML 里主题属性
    优先于显式 w:ascii / w:eastAsia，python-docx 只加不删，不清就等于没设。
    """
    rPr = _rpr(obj)
    if latin or cjk:
        rf = rPr.get_or_add_rFonts()
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rf.attrib.pop(qn("w:" + attr), None)
        if latin:
            for attr in ("ascii", "hAnsi", "cs"):
                rf.set(qn("w:" + attr), latin)
        if cjk:
            rf.set(qn("w:eastAsia"), cjk)
            rf.set(qn("w:hint"), "eastAsia")   # 中英混排时标点归位
    if size is not None:
        for tag in ("w:sz", "w:szCs"):
            for old in rPr.findall(qn(tag)):
                rPr.remove(old)
        rPr.append(_el("w:sz", val=int(round(size * 2))))
        rPr.append(_el("w:szCs", val=int(round(size * 2))))
    if bold is not None:
        for old in rPr.findall(qn("w:b")):
            rPr.remove(old)
        rPr.append(_el("w:b", val="1" if bold else "0"))
    if italic is not None:
        for old in rPr.findall(qn("w:i")):
            rPr.remove(old)
        rPr.append(_el("w:i", val="1" if italic else "0"))
    if color is not None:
        for old in rPr.findall(qn("w:color")):
            rPr.remove(old)
        rPr.append(_el("w:color", val=color.lstrip("#")))
    return obj


def set_first_line_indent_chars(target, chars=2):
    """中文首行缩进“2 字符”（w:firstLineChars），字号变了也仍是 2 个字。"""
    pf = target.paragraph_format if hasattr(target, "paragraph_format") else target
    el = pf.element if hasattr(pf, "element") else pf._element
    pPr = el.get_or_add_pPr() if hasattr(el, "get_or_add_pPr") else el
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.insert_element_before(
            ind, "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
            "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
        )
    ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
    ind.set(qn("w:firstLine"), "0")   # 由 Chars 接管
    return target


# --- 页面与主题 -------------------------------------------------------------
# 竖版 / 横版的默认边距（上, 右, 下, 左），单位 cm
PORTRAIT_MARGINS = (2.54, 3.18, 2.54, 3.18)
LANDSCAPE_MARGINS = (2.0, 2.8, 2.0, 2.8)


def _page_size(size, orientation):
    w, h = (Cm(21), Cm(29.7)) if size.upper() == "A4" else (Inches(8.5), Inches(11))
    return (h, w) if str(orientation).lower().startswith("land") else (w, h)


def body_width_cm(section):
    """版心宽度（cm）。分栏时返回单栏宽——行长要按栏宽算，不是按页宽。"""
    total = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    cols = section._sectPr.find(qn("w:cols"))
    n = int(cols.get(qn("w:num")) or 1) if cols is not None else 1
    if n > 1:
        space = int(cols.get(qn("w:space")) or 425) / 567.0   # twips → cm
        total = (total - space * (n - 1)) / n
    return total


def set_columns(section, num=2, space_cm=0.75, line_between=False):
    """分栏。横版正文必须分栏或收窄版心，否则一行会到 55+ 字，没法读。"""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = _el("w:cols", num=num, space=int(round(space_cm * 567)))
    if line_between:
        cols.set(qn("w:sep"), "1")
    sectPr.insert_element_before(
        cols, "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
        "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid", "w:printerSettings",
    )
    return section


def _apply_section_page(sec, size, orientation, margins, grid_lines=None):
    w, h = _page_size(size, orientation)
    landscape = str(orientation).lower().startswith("land")
    sec.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = w, h      # 只设 orientation 不换宽高是无效的
    top, right, bottom, left = margins or (
        LANDSCAPE_MARGINS if landscape else PORTRAIT_MARGINS)
    sec.top_margin, sec.bottom_margin = Cm(top), Cm(bottom)
    sec.left_margin, sec.right_margin = Cm(left), Cm(right)
    sec.header_distance, sec.footer_distance = Cm(1.25 if landscape else 1.5), Cm(1.25 if landscape else 1.75)
    if grid_lines:
        usable = sec.page_height - sec.top_margin - sec.bottom_margin
        pitch = int(round(Emu(usable).twips / grid_lines))
        for old in sec._sectPr.findall(qn("w:docGrid")):
            sec._sectPr.remove(old)
        sec._sectPr.append(_el("w:docGrid", type="lines", linePitch=pitch))
    return sec


def setup_page(doc, size="A4", margins=None, grid_lines=None, orientation="portrait"):
    """把所有节设成 A4（或 Letter）+ 中文习惯边距。

    python-docx 默认是 US Letter（21.59×27.94），SKILL.md 的字号/行长推导
    全部基于 A4，不设这一步整篇版心就是错的。
    margins = (上, 右, 下, 左)，单位 cm；orientation = 'portrait' | 'landscape'。
    """
    w, h = _page_size(size, orientation)
    landscape = str(orientation).lower().startswith("land")
    top, right, bottom, left = margins or (
        LANDSCAPE_MARGINS if landscape else PORTRAIT_MARGINS)
    for sec in doc.sections:
        sec.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = w, h
        sec.top_margin, sec.bottom_margin = Cm(top), Cm(bottom)
        sec.left_margin, sec.right_margin = Cm(left), Cm(right)
        sec.header_distance, sec.footer_distance = Cm(1.5), Cm(1.75)
        if grid_lines:
            usable = sec.page_height - sec.top_margin - sec.bottom_margin
            pitch = int(round(Emu(usable).twips / grid_lines))
            for old in sec._sectPr.findall(qn("w:docGrid")):
                sec._sectPr.remove(old)
            sec._sectPr.append(_el("w:docGrid", type="lines", linePitch=pitch))
    return doc


def apply_theme(doc, theme="navy", preset="business"):
    """一次性定义 Normal / Heading 1-3 / Caption，全文生效。

    样式层定义 > run 层硬编码：改一处全文跟着改，也不会漏掉 eastAsia。
    只有“破例”的文字才在 run 上调 set_font()。
    """
    t = THEMES[theme] if isinstance(theme, str) else theme
    p = dict(PRESETS[preset]) if isinstance(preset, str) else dict(preset)
    accent = t["accent"]

    # Normal：正文字体、行距、段后距、首行缩进
    normal = doc.styles["Normal"]
    set_font(normal, latin=p["latin"], cjk=p["cjk"], size=p["body"], color=INK)
    pf = normal.paragraph_format
    if p.get("line_exact"):
        pf.line_spacing = Pt(p["line_exact"])
    elif p.get("line"):
        pf.line_spacing = p["line"]
    pf.space_after = Pt(p.get("space_after", 8))
    pf.space_before = Pt(0)
    pf.widow_control = True
    if p.get("indent_chars"):
        set_first_line_indent_chars(normal, p["indent_chars"])

    # 标题：段前 > 段后（标题贴近它管辖的内容），层级差 ≥2pt 且换字重
    for lvl, before, after in ((1, 18, 8), (2, 14, 6), (3, 12, 6)):
        st = doc.styles[f"Heading {lvl}"]
        set_font(st, latin=p.get("h_latin", p["latin"]), cjk=p.get("h_cjk", p["cjk"]),
                 size=p[f"h{lvl}"], bold=True,
                 color=accent if lvl == 1 else (accent if lvl == 2 else INK))
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.25
        set_first_line_indent_chars(st, 0)

    # 图注/表注：比正文小一档、灰字、居中、不缩进
    try:
        cap = doc.styles["Caption"]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_font(cap, latin=p["latin"], cjk=p["cjk"],
             size=max(8, p["body"] - 3), color=MUTED, bold=False, italic=False)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    set_first_line_indent_chars(cap, 0)

    doc._theme = t          # 供后续 helper 取色
    doc._preset = p
    return doc


def new_document(theme="navy", preset="business", size="A4",
                 orientation="portrait", columns=1):
    """建一个已经设好页面 + 字体 + 样式的空文档。正常情况下从这里开始。

    orientation='landscape' 时若不分栏，版心会宽到 24cm（一行 55+ 字），
    所以横版正文建议同时给 columns=2。表格/图为主的横版可以不分栏。
    """
    doc = Document()
    p = PRESETS[preset] if isinstance(preset, str) else preset
    setup_page(doc, size=size, margins=p.get("margins"),
               grid_lines=p.get("grid_lines"), orientation=orientation)
    apply_theme(doc, theme=theme, preset=preset)
    if columns > 1:
        set_columns(doc.sections[0], columns)
    return doc


def _theme_of(doc):
    return getattr(doc, "_theme", THEMES["navy"])


# --- 域（页码 / 目录） ------------------------------------------------------
def add_field(paragraph, instr, cached=None, dirty=True):
    """插入一个 Word 域。cached 是渲染成 PDF 时显示的占位内容。

    没有 cached 的域，LibreOffice 转 PDF 时是空白——目录尤其明显。
    """
    r = paragraph.add_run()
    begin = _el("w:fldChar", fldCharType="begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    r._r.append(begin)
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r._r.append(instr_el)
    r._r.append(_el("w:fldChar", fldCharType="separate"))
    if cached is not None:
        paragraph.add_run(cached)
    end = paragraph.add_run()
    end._r.append(_el("w:fldChar", fldCharType="end"))
    return paragraph


def add_page_number(paragraph, fmt="cn", total=True):
    """页脚页码。fmt='cn' → 第 X 页 共 Y 页；fmt='plain' → X。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fmt == "cn":
        paragraph.add_run("第 ")
        add_field(paragraph, "PAGE", cached="1")
        paragraph.add_run(" 页")
        if total:
            paragraph.add_run("，共 ")
            add_field(paragraph, "NUMPAGES", cached="1")
            paragraph.add_run(" 页")
    else:
        add_field(paragraph, "PAGE", cached="1")
    for run in paragraph.runs:
        set_font(run, size=9, color=MUTED)
    return paragraph


def add_toc(doc, title="目录", levels=(1, 3), page_break=True):
    """插入自动目录域。

    诚实说明：域里没有页码缓存，**LibreOffice 转出的 PDF 里目录是提示文字，
    不是真目录**；在 Word 中打开会提示更新（w:dirty），或 Ctrl+A 后按 F9。
    交付前要把这句话告诉用户。需要 PDF 里也有真目录时，只能手工写静态目录。
    """
    if title:
        h = doc.add_paragraph(title)
        try:
            h.style = doc.styles["TOC Heading"]
        except KeyError:
            h.style = doc.styles["Heading 1"]
        p_ = dict(getattr(doc, "_preset", PRESETS["business"]))
        set_font(h.runs[0], latin=p_.get("h_latin"), cjk=p_.get("h_cjk"),
                 size=p_["h1"], bold=True, color=_theme_of(doc)["accent"])
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    add_field(
        p,
        r'TOC \o "%d-%d" \h \z \u' % levels,
        cached="（本目录为 Word 自动目录域：在 Word 中打开后按 Ctrl+A、F9 更新即可生成带页码的目录）",
    )
    for run in p.runs:
        set_font(run, size=10, color=MUTED)
    if page_break:
        page_break_section(doc)
    return doc


def page_break_section(doc):
    """换页而不留空段落。

    用分节符换页，比“空段落 + 手工分页符”干净：不留空段落（检查器不会告警，
    也不会在预览里显示成一个方块），方向和页码设置都自动延续。
    """
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.header.is_linked_to_previous = True
    sec.footer.is_linked_to_previous = True
    for old in sec._sectPr.findall(qn("w:pgNumType")):
        sec._sectPr.remove(old)          # 否则会继承“页码从 1 开始”，翻页又归零
    return sec


# --- 封面与分节 -------------------------------------------------------------
def add_rule(paragraph, color=None, width_pt=1.5, space=6):
    """段落下方的一条细线——封面点缀、章节分隔用，比色块克制。"""
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    bdr = OxmlElement("w:pBdr")
    bdr.append(_el("w:bottom", val="single", sz=int(width_pt * 8),
                   space=space, color=(color or "1F4E79").lstrip("#")))
    pPr.insert_element_before(
        bdr, "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
        "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
        "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
        "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
        "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
        "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
    )
    return paragraph


def add_cover(doc, title, subtitle=None, meta=None, page_break=True):
    """封面：标题落在上三分处，不用空行排版，一条细线点缀。

    meta 是底部的单位/日期等，list[str]。
    """
    t = _theme_of(doc)
    p_ = dict(getattr(doc, "_preset", PRESETS["business"]))

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(150)   # 用段前距造留白，不用空行
    tp.paragraph_format.space_after = Pt(10)
    set_first_line_indent_chars(tp, 0)
    set_font(tp.add_run(title), latin=p_.get("h_latin"), cjk=p_.get("h_cjk"),
             size=30, bold=True, color=INK)

    rule_p = doc.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_p.paragraph_format.space_after = Pt(14)
    add_rule(rule_p, color=t["accent"], width_pt=1.5)

    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(0)
        set_first_line_indent_chars(sp, 0)
        set_font(sp.add_run(subtitle), latin=p_.get("h_latin"), cjk=p_.get("h_cjk"),
                 size=15, bold=False, color=MUTED)

    for i, line in enumerate(meta or []):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(180 if i == 0 else 0)
        mp.paragraph_format.space_after = Pt(4)
        set_first_line_indent_chars(mp, 0)
        set_font(mp.add_run(line), latin=p_["latin"], cjk=p_["cjk"],
                 size=12, color=MUTED)

    if page_break:
        page_break_section(doc)
    return doc


def _last_section_empty(doc):
    """最后一节里还没有任何内容——可以直接复用，不必再插一个分节符。"""
    seen_last_break = False
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if seen_last_break and (child.xpath("string(.)").strip()
                                    or child.find(qn("w:r")) is not None):
                return False
            if child.find(f"{qn('w:pPr')}/{qn('w:sectPr')}") is not None:
                seen_last_break = True
        elif child.tag == qn("w:tbl") and seen_last_break:
            return False
    return seen_last_break


def start_body_section(doc, page_number=True, header_text=None, restart=True):
    """正文另起一节：封面不带页码，正文页码从 1 开始。

    紧跟 add_cover() 调用时会复用封面留下的空节，不会多出一张白页。
    """
    sec = doc.sections[-1] if _last_section_empty(doc) else doc.add_section(WD_SECTION.NEW_PAGE)
    sec.footer.is_linked_to_previous = False
    sec.header.is_linked_to_previous = False
    if restart:
        for old in sec._sectPr.findall(qn("w:pgNumType")):
            sec._sectPr.remove(old)
        pg = _el("w:pgNumType", start=1)
        sec._sectPr.insert_element_before(
            pg, "w:cols", "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
            "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid", "w:printerSettings",
        )
    if page_number:
        add_page_number(sec.footer.paragraphs[0])
    if header_text:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(hp.add_run(header_text), size=9, color=MUTED)
        add_rule(hp, color=_theme_of(doc)["accent"], width_pt=0.5, space=2)
    return sec


def _new_oriented_section(doc, orientation, margins=None, columns=1,
                          keep_footer=True):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    prev = doc.sections[-2] if len(doc.sections) > 1 else None
    size = "A4" if round(max(sec.page_width.cm, sec.page_height.cm), 1) == 29.7 else "Letter"
    _apply_section_page(sec, size, orientation, margins)
    for old in sec._sectPr.findall(qn("w:cols")):
        sec._sectPr.remove(old)
    if columns > 1:
        set_columns(sec, columns)
    if keep_footer and prev is not None:
        sec.footer.is_linked_to_previous = True     # 页码延续，不重编
        sec.header.is_linked_to_previous = True
    return sec


def add_landscape_section(doc, columns=1, margins=None):
    """插入一个横版节——**局部例外**，用于确实无法收窄的宽表 / 宽图 / 大图表。

        add_landscape_section(doc)
        t = doc.add_table(rows=..., cols=9); style_table(t, doc=doc)
        add_portrait_section(doc)          # 切回竖版继续写正文

    先想清楚：页面方向由**文档内容性质**决定（连续文字要读 → 竖版；
    并排比较、看板、图表要看 → 横版），是动笔前的整体决定。
    一张表放不下时优先改表（转置 / 合并列 / 拆表 / 降一档字号），
    这个函数是这些手段都用尽之后的最后一步。
    """
    return _new_oriented_section(doc, "landscape", margins, columns)


def add_portrait_section(doc, columns=1, margins=None):
    """切回竖版。横版节之后必须调用它，否则后面全变横版。"""
    return _new_oriented_section(doc, "portrait", margins, columns)


# --- 表格 -------------------------------------------------------------------
def auto_column_widths(table, avail_cm, min_col_cm=1.6, damp=0.65, font_pt=9):
    """按各列的内容量分配列宽，合计正好铺满版心。

    直接按字数正比分配会让长文本列吃掉一切，所以用 volume**0.65 压一下。
    每列的最小宽度按**该列最长的短标签**算，不是一个统一值——“MaaS平台”“模型安全”
    这类 4-6 字的分类标签被压到 1.6cm 就会折行，很难看，而联系人列 2cm 就够。
    """
    n = len(table.columns)
    if n == 0:
        return []

    def w_cm(text):
        return sum(1.0 if CJK_RE.search(ch) else 0.5 for ch in text) * font_pt * 2.54 / 72

    vols, floors = [], []
    for c in range(n):
        chars, token_w, label_w = 0.0, 0.0, 0.0
        for r in range(len(table.rows)):
            text = table.rows[r].cells[c].text
            chars += w_cm(text) * 72 / 2.54 / font_pt      # 折回“字身数”
            for tok in text.split():
                # 中文可以在任意字之间断行，数字/英文串不能——手机号被折成两行最难看
                if tok and not CJK_RE.search(tok):
                    token_w = max(token_w, w_cm(tok))
            if 0 < len(text) <= 8:                          # 短标签也不该折行
                label_w = max(label_w, w_cm(text))
        vols.append(max(chars / max(1, len(table.rows)), 1.0))
        # 留白比理论值多给一点：实际渲染的字体常比计算用的字身宽（跨平台字体替换尤其明显），
        # 卡着算会让手机号、四字标签正好折成两行。上限防止某一列被一个超长 token 绑架。
        floors.append(min(max(min_col_cm, token_w + 0.4, label_w + 0.5), avail_cm * 0.15))

    weights = [v ** damp for v in vols]
    if sum(floors) >= avail_cm:              # 下限本身就超宽：按比例缩，保住相对关系
        scale = avail_cm / sum(floors)
        return [f * scale for f in floors]

    # 触到下限的列要**锁住**——不锁的话下一轮它又回到自由分配里，来回震荡，
    # 结果比不设下限还窄（实测踩过：抬高下限反而把列压到了 1.07cm）。
    widths = [0.0] * n
    locked = {}
    while True:
        free = [i for i in range(n) if i not in locked]
        spare = avail_cm - sum(locked.values())
        if not free:
            break
        wsum = sum(weights[i] for i in free) or 1
        newly = [i for i in free if spare * weights[i] / wsum < floors[i]]
        if not newly:
            for i in free:
                widths[i] = spare * weights[i] / wsum
            break
        for i in newly:
            locked[i] = floors[i]
    for i, v in locked.items():
        widths[i] = v
    return widths


def fit_table_to_page(table, doc=None, section=None, col_widths_cm=None,
                      min_col_cm=1.6, font_pt=None):
    """把表格铺满版心，并按内容量分配列宽。

    **为什么必须显式设列宽**：Word 默认是 autofit，会**忽略**你设的 cell.width
    按内容自行收缩——实测里表格因此只占了半页宽，右边一大片空白，而声明的列宽
    看起来完全正确。所以这里同时做三件事：`w:tblLayout=fixed`、`w:tblW=100%`、
    把宽度写进 tblGrid 和每一个单元格。三样缺一，Word 就会自作主张。

    col_widths_cm 显式给定时按它排（会等比缩放到版心内），否则按内容量自动分配。
    """
    sec = section or (doc.sections[-1] if doc is not None else None)
    avail = body_width_cm(sec) if sec is not None else 15.0
    tbl = table._tbl
    tblPr = tbl.tblPr

    if font_pt is None:
        preset = dict(getattr(doc, "_preset", PRESETS["business"])) if doc else PRESETS["business"]
        font_pt = max(9, preset["body"] - 1.5)
    widths = list(col_widths_cm) if col_widths_cm else auto_column_widths(
        table, avail, min_col_cm, font_pt=font_pt)
    if col_widths_cm:
        total = sum(widths) or 1
        widths = [w * avail / total for w in widths]      # 归一化到版心

    for tag in ("w:tblW", "w:tblLayout"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    tblPr.insert_element_before(_el("w:tblW", w=5000, type="pct"),
                                "w:jc", "w:tblCellSpacing", "w:tblInd",
                                "w:tblBorders", "w:shd", "w:tblLayout",
                                "w:tblCellMar", "w:tblLook")
    tblPr.insert_element_before(_el("w:tblLayout", type="fixed"),
                                "w:tblCellMar", "w:tblLook")
    table.autofit = False

    grid = tbl.find(qn("w:tblGrid"))          # tblGrid 与单元格宽度都要写
    if grid is not None:
        for i, col in enumerate(grid.findall(qn("w:gridCol"))):
            if i < len(widths):
                col.set(qn("w:w"), str(int(round(widths[i] * 567))))
    for i, w in enumerate(widths):
        if i >= len(table.columns):
            break
        table.columns[i].width = Cm(w)
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    return table


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = _el("w:shd", val="clear", color="auto", fill=hex_color.lstrip("#"))
    tcPr.insert_element_before(
        shd, "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
        "w:vAlign", "w:hideMark",
    )
    return cell


def _set_table_borders(table, inner=LINE, outer=LINE, inner_pt=0.5, outer_pt=1.0):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, pt, color in (
        ("top", outer_pt, outer), ("left", outer_pt, outer),
        ("bottom", outer_pt, outer), ("right", outer_pt, outer),
        ("insideH", inner_pt, inner), ("insideV", inner_pt, inner),
    ):
        borders.append(_el("w:" + edge, val="single", sz=max(2, int(pt * 8)),
                           space=0, color=color.lstrip("#")))
    tblPr.insert_element_before(borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")


def _set_cell_margins(table, top=60, left=108, bottom=60, right=108):
    """单元格内边距，单位 dxa（1/20 pt）。Word 默认上下为 0，太挤。"""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(old)
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.insert_element_before(mar, "w:tblLook")


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:tblHeader")):
        trPr.remove(old)
    trPr.append(OxmlElement("w:tblHeader"))


def style_table(table, doc=None, theme=None, header=True, zebra=None,
                font_size=None, align_right_cols=(), header_style="accent",
                fit=True, col_widths_cm=None):
    """把一张裸表变成可交付的表：细浅灰内框 + 着色表头 + 内边距 + 斑马纹 + 跨页表头。

    header_style: 'accent' = 强调色底白字；'tint' = 浅底深字（更克制，适合密表）。
    align_right_cols: 数字列索引，右对齐。
    """
    t = theme or _theme_of(doc) if (theme or doc) else THEMES["navy"]
    if isinstance(t, str):
        t = THEMES[t]
    preset = dict(getattr(doc, "_preset", PRESETS["business"])) if doc else PRESETS["business"]
    size = font_size or max(9, preset["body"] - 1.5)
    rows = len(table.rows)
    if zebra is None:
        zebra = rows > 7

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if fit:
        fit_table_to_page(table, doc=doc, col_widths_cm=col_widths_cm, font_pt=size)
    # 长文本表格垂直居中会让每行的文字浮在中间、参差不齐，超过一行就该顶对齐
    longest = max((len(c.text) for row in table.rows for c in row.cells), default=0)
    body_valign = "top" if longest > 40 else "center"
    if "Table Grid" in (table.style.name or ""):
        table.style = table.style          # 保留，但下面的边框会覆盖掉黑网格
    _set_table_borders(table, inner=LINE, outer=t["accent"] if header else LINE,
                       inner_pt=0.5, outer_pt=0.75)
    _set_cell_margins(table)

    for r, row in enumerate(table.rows):
        is_header = header and r == 0
        for c, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:vAlign")):
                tcPr.remove(old)
            v = _el("w:vAlign", val="center" if is_header else body_valign)
            tcPr.insert_element_before(v, "w:hideMark")
            if is_header:
                shade_cell(cell, t["accent"] if header_style == "accent" else t["tint"])
            elif zebra and (r % 2 == 0):
                shade_cell(cell, ZEBRA)
            for para in cell.paragraphs:
                set_first_line_indent_chars(para, 0)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.15
                if is_header:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif c in align_right_cols:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in para.runs:
                    set_font(
                        run, latin=preset["latin"], cjk=preset["cjk"], size=size,
                        bold=is_header,
                        color=("FFFFFF" if header_style == "accent" else t["accent"])
                        if is_header else INK,
                    )
    if header and rows:
        _repeat_header(table.rows[0])
    return table


def set_column_alignment(table, col, align="right", skip_header=True):
    a = {"right": WD_ALIGN_PARAGRAPH.RIGHT, "left": WD_ALIGN_PARAGRAPH.LEFT,
         "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i, row in enumerate(table.rows):
        if skip_header and i == 0:
            continue
        for para in row.cells[col].paragraphs:
            para.alignment = a
    return table


# --- 图 ---------------------------------------------------------------------
def add_figure(doc, image_path, caption=None, width_in=5.8, source=None):
    """插图 + 灰字图注（含数据来源）。图注用 Caption 样式，不手工设格式。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(p, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(image_path, width=Inches(width_in))
    if caption or source:
        cp = doc.add_paragraph(style=doc.styles["Caption"])
        text = caption or ""
        if source:
            text = f"{text}（数据来源：{source}）" if text else f"数据来源：{source}"
        cp.add_run(text)
        set_font(cp.runs[0], size=max(8, _preset_body(doc) - 3), color=MUTED)
    return p


def _preset_body(doc):
    return dict(getattr(doc, "_preset", PRESETS["business"]))["body"]
