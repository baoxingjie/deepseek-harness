"""内容保真比对：成品 DOCX 与源文件（xlsx/csv）逐格对照。

    python scripts/verify_source.py 源.xlsx 成品.docx
    python scripts/verify_source.py 源.xlsx 成品.docx --sheet Sheet1 --json

转换类任务（Excel→Word、台账、周报）**不允许改写、缩写、删减原文**，
但“我看着挺像”是靠不住的——本技能自己就栽过：只比了两列就宣布“全部一致”，
实际表头少了一个字、标题被拆写、两级任务列被拍平、单元格内换行被换成了空格。

所以这个脚本做的是**双向全量比对**：
  · 源里有、成品里找不到的  → 漏内容或被改写（最严重）
  · 成品里有、源里没有的    → 新增或改写的文字（标题、页眉这类合法补充需要人工确认）
  · 只差空白/换行的         → 单独列出，别混进“内容一致”里

退出码：0 = 无内容差异；1 = 有内容差异。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from os.path import splitext

# Windows 控制台默认 GBK，编不出的字符会让 print 抛 UnicodeEncodeError 中断脚本。
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(errors="replace")

WS = re.compile(r"\s+")


def norm(s):
    return (s or "").strip()


def loose(s):
    """去掉所有空白后的形态——用来区分“内容不同”与“只是换行/空格不同”。"""
    return WS.sub("", s or "")


def read_source(path, sheet=None):
    """返回源文件里所有非空单元格的文本。合并单元格只有左上角有值，天然只取一次。"""
    ext = splitext(path)[1].lower()
    cells = []
    if ext in (".xlsx", ".xlsm"):
        try:
            import openpyxl
        except ImportError:
            raise SystemExit("需要 openpyxl：pip install openpyxl")
        wb = openpyxl.load_workbook(path, data_only=True)
        sheets = [wb[sheet]] if sheet else wb.worksheets
        for ws in sheets:
            for row in ws.iter_rows():
                for c in row:
                    if c.value is not None and norm(str(c.value)):
                        cells.append((f"{ws.title}!{c.coordinate}", norm(str(c.value))))
    elif ext in (".csv", ".tsv"):
        import csv
        delim = "\t" if ext == ".tsv" else ","
        with open(path, newline="", encoding="utf-8-sig") as f:
            for ri, row in enumerate(csv.reader(f, delimiter=delim), 1):
                for ci, v in enumerate(row):
                    if norm(v):
                        cells.append((f"R{ri}C{ci+1}", norm(v)))
    else:
        raise SystemExit(f"不支持的源文件类型：{ext}（支持 xlsx/xlsm/csv/tsv）")
    return cells


def _para_text(p_el, W):
    """一个 w:p 的文字，w:br / w:tab 还原成 \\n 与制表符。"""
    parts = []
    for node in p_el.iter():
        if node.tag == W + "t":
            parts.append(node.text or "")
        elif node.tag == W + "br":
            parts.append("\n")
        elif node.tag == W + "tab":
            parts.append("\t")
    return "".join(parts)


def read_docx(path):
    from docx import Document
    doc = Document(path)
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    out = []
    for i, p in enumerate(doc.paragraphs):
        if norm(p.text):
            out.append((f"段落#{i}", norm(p.text)))
    # 直接遍历 XML 里的 w:tc：合并单元格在 XML 中只出现一次，不必去重
    # （用 id(cell._tc) 去重会踩 CPython 的 id 复用，实测把 45 个单元格读成了 8 个）
    for ti, t in enumerate(doc.tables):
        for ri, tr in enumerate(t._tbl.findall(W + "tr")):
            for ci, tc in enumerate(tr.findall(W + "tc")):
                text = norm("\n".join(_para_text(p, W) for p in tc.findall(W + "p")))
                if text:
                    out.append((f"表{ti}[{ri},{ci}]", text))
    for si, s in enumerate(doc.sections):
        for part, name in ((s.header, "页眉"), (s.footer, "页脚")):
            for p in part.paragraphs:
                if norm(p.text):
                    out.append((f"{name}#{si}", norm(p.text)))
    return out


def compare(src, dst):
    dst_exact = {}
    dst_loose = {}
    for where, text in dst:
        dst_exact.setdefault(text, []).append(where)
        dst_loose.setdefault(loose(text), []).append(where)

    missing, whitespace_only = [], []
    matched_dst = set()
    for where, text in src:
        if text in dst_exact:
            matched_dst.update(dst_exact[text])
            continue
        if loose(text) in dst_loose:
            whitespace_only.append((where, text, dst_loose[loose(text)][0]))
            matched_dst.update(dst_loose[loose(text)])
            continue
        # 被拆写/被并入更长的文本里，也算改动，但单独标注更好定位
        holder = next((w for w, t in dst if loose(text) and loose(text) in loose(t)), None)
        missing.append((where, text, holder))

    src_loose = {loose(t) for _, t in src}
    extra = [(where, text) for where, text in dst
             if where not in matched_dst and loose(text) not in src_loose
             and not any(loose(text) in loose(t) for _, t in src)]
    return missing, whitespace_only, extra


def main():
    ap = argparse.ArgumentParser(description="DOCX 与源文件的内容保真比对")
    ap.add_argument("source")
    ap.add_argument("docx")
    ap.add_argument("--sheet", default=None)
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--allow-extra", action="store_true",
                    help="允许成品里有源文件没有的文字（页码、页眉这类）")
    args = ap.parse_args()

    src = read_source(args.source, args.sheet)
    dst = read_docx(args.docx)
    missing, ws_only, extra = compare(src, dst)

    if args.json:
        print(json.dumps({"missing": missing, "whitespace_only": ws_only, "extra": extra},
                         ensure_ascii=False, indent=2))
    else:
        print(f"源 {len(src)} 个非空单元格 → 成品 {len(dst)} 段文本\n")
        if missing:
            print(f"× 源里有、成品里找不到（{len(missing)} 处）——漏内容或被改写：")
            for where, text, holder in missing:
                tip = f"    疑似被并入：{holder}" if holder else ""
                print(f"  [{where}] {text[:60]}{'…' if len(text) > 60 else ''}{tip}")
        if ws_only:
            print(f"\n! 只有空白/换行不同（{len(ws_only)} 处）——"
                  "单元格内的换行是原文的一部分，别替换成空格：")
            for where, text, at in ws_only:
                print(f"  [{where}] → [{at}] {text[:50]}")
        if extra:
            print(f"\n! 成品里有、源里没有（{len(extra)} 处）——确认是合法补充还是改写：")
            for where, text in extra:
                print(f"  [{where}] {text[:60]}")
        if not (missing or ws_only or extra):
            print("√ 内容完全一致（含表头、标题、单元格内换行）。")

    bad = bool(missing) or bool(ws_only) or (bool(extra) and not args.allow_extra)
    sys.exit(1 if bad else 0)


if __name__ == "__main__":
    main()
