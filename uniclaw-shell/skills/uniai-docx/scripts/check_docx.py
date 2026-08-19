"""DOCX 交付前自检：把 SKILL.md 的硬规则跑成一道闸门。

    python scripts/check_docx.py out/report.docx
    python scripts/check_docx.py out/report.docx --json      # 机器可读
    python scripts/check_docx.py out/report.docx --letter    # 允许 US Letter 页面

有 ERROR 就是返工（退出码 1）；WARN 需要你看一眼再决定。
它检查的是“抽文字看不出、但渲染出来一眼就丑”的那些问题：缺东亚字体、
假标题、空行排版、裸表头、全黑网格、多强调色、行长失控、占位符残留，
以及乱码、直引号、半角标点这类字符层面的脏东西。

它不能替代视觉复检——分页、溢出、图片对齐仍然要看 render_docx.py 的 PNG。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter

from docx import Document
from docx.oxml.ns import qn

# Windows 控制台默认是 GBK，编不出的字符会让 print 抛 UnicodeEncodeError 中断整个脚本。
# 输出退化成 ? 也比崩掉强；同理，下面的提示图标只用 GBK 里有的字符。
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(errors="replace")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
CJK = re.compile(r"[㐀-鿿豈-﫿　-〿＀-￯]")
PLACEHOLDER = re.compile(r"(XXX|xxx|TODO|TBD|待补充|待填写|占位|Lorem ipsum|\{\{.*?\}\}|\[.{0,12}?待.{0,6}?\])")
NEUTRAL = {"000000", "1F2937", "595959", "auto", "D9D9D9", "F7F8FA", "FFFFFF",
           "F2F2F2", "808080", "404040", "262626", "7F7F7F", "BFBFBF", "E7E6E6"}
END_PUNCT = "。！？；：,.!?;:、"
THEME_ATTRS = ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")

# --- 字符层面的问题：抽文字能看出来，但只看渲染图容易滑过去 -------------------
# 直引号：中文里应当用 “ ”‘ ’。ASCII 直引号在 Word 里既不成对也不会自动弯，
# 而且它正是生成脚本 SyntaxError 的同一个源头（见 SKILL.md「生成脚本别被字符卡死」）。
STRAIGHT_QUOTE = re.compile(r"""(?:(?<=[㐀-鿿])["']|["'](?=[㐀-鿿]))""")
# 半角标点夹在中文之间：中文正文该用全角 ，；！？。句点不查——小数点、网址、英文缩写太多。
HALF_PUNCT = re.compile(r"(?:(?<=[㐀-鿿])[,;!?]|[,;!?](?=[㐀-鿿]))")
# 乱码痕迹：U+FFFD 替换字符、文中残留的 BOM、GBK/UTF-8 互串的经典产物。
MOJIBAKE = re.compile("[\ufffd\ufeff]|锟斤拷|鎴戜滑")

# 版面估算的两个系数，用真实样张校准过（见 selftest 的 page_estimate 用例）。
# LINE_H：一行占多少个字身。Word 单倍行距约 1.2，本技能表格默认 1.15，取 1.25 略留余量。
# WRAP_BIAS：折行数取整偏置。原来是 0.45，几乎总是向上进位，导致长文本表格高估约 40%——
# 实测 8pt 全文版只占 2/3 页，却被估成 2 页，会把智能体推去做没必要的缩字号。
LINE_H = 1.25
WRAP_BIAS = 0.15
# 纯几何推导仍比实际偏高约 15%（真实文本不会每行都排满，折行点还会被标点挤走），
# 这个系数是拿渲染实测的样张标定的：8pt 全文周报实测占 0.68 页，未校准估算 0.80。
CALIBRATION = 0.85


class Report:
    def __init__(self):
        self.items = []

    def add(self, level, code, msg, where=""):
        self.items.append({"level": level, "code": code, "msg": msg, "where": where})

    error = lambda self, c, m, w="": self.add("ERROR", c, m, w)
    warn = lambda self, c, m, w="": self.add("WARN", c, m, w)
    info = lambda self, c, m, w="": self.add("INFO", c, m, w)


def snippet(text, n=28):
    text = " ".join((text or "").split())
    return (text[:n] + "…") if len(text) > n else text


# --- 样式链：某个 styleId 最终有没有 eastAsia 字体 ---------------------------
def style_font_map(doc):
    out = {}
    for st in doc.styles.element.findall(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        rPr = st.find(qn("w:rPr"))
        rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
        based = st.find(qn("w:basedOn"))
        out[sid] = {
            "eastAsia": rf.get(qn("w:eastAsia")) if rf is not None else None,
            "theme": [a for a in THEME_ATTRS if rf is not None and rf.get(qn("w:" + a))],
            "basedOn": based.get(qn("w:val")) if based is not None else None,
            "name": (st.find(qn("w:name")).get(qn("w:val"))
                     if st.find(qn("w:name")) is not None else sid),
        }
    return out


def resolved_eastasia(sid, smap, depth=0):
    if not sid or sid not in smap or depth > 8:
        return None
    node = smap[sid]
    return node["eastAsia"] or resolved_eastasia(node["basedOn"], smap, depth + 1)


def docdefaults_eastasia(doc):
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is None:
        return None
    rf = dd.find(f"{W}rPrDefault/{W}rPr/{W}rFonts")
    return rf.get(qn("w:eastAsia")) if rf is not None else None


# --- 各项检查 ---------------------------------------------------------------
def columns_of(sec):
    cols = sec._sectPr.find(qn("w:cols"))
    return int(cols.get(qn("w:num")) or 1) if cols is not None else 1


def body_width(sec):
    """版心宽（cm）；分栏时返回单栏宽——行长按栏宽算。"""
    total = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    n = columns_of(sec)
    if n > 1:
        cols = sec._sectPr.find(qn("w:cols"))
        space = int(cols.get(qn("w:space")) or 425) / 567.0
        total = (total - space * (n - 1)) / n
    return total


def is_landscape(sec):
    return sec.page_width.cm > sec.page_height.cm


def check_page(doc, rep, allow_letter):
    """页面尺寸与方向。横版是合法选择，但横版正文必须分栏或收窄版心。"""
    for i, sec in enumerate(doc.sections):
        w_cm, h_cm = round(sec.page_width.cm, 1), round(sec.page_height.cm, 1)
        a4 = sorted((w_cm, h_cm)) == [21.0, 29.7]
        if not a4 and not allow_letter:
            rep.error("PAGE_SIZE",
                      f"第 {i+1} 节页面是 {w_cm}×{h_cm}cm，不是 A4（21×29.7 或横版 29.7×21）。"
                      "python-docx 默认 US Letter，必须显式 setup_page()。", f"sec#{i}")
        # 横版宽版心只有在“这一节真的有大段连续正文”时才是问题；
        # 只放宽表/宽图的横版节是合理用法，不该报警。
        prose = sum(len(b.text) for si, b in walk_body(doc)
                    if si == i and hasattr(b, "runs")
                    and not (b.style.name or "").startswith(("Heading", "Caption", "TOC")))
        if is_landscape(sec) and columns_of(sec) == 1 and body_width(sec) > 20 and prose > 200:
            rep.warn("LANDSCAPE_WIDE",
                     f"第 {i+1} 节是横版、未分栏，版心 {body_width(sec):.1f}cm，"
                     f"却有约 {prose} 字连续正文——一行会到 50+ 字，读起来很累。"
                     "横版正文应分两栏（set_columns）或收窄版心。", f"sec#{i}")

    orients = [("横" if is_landscape(s) else "竖") for s in doc.sections]
    if len(set(orients)) > 1:
        rep.info("MIXED_ORIENTATION",
                 f"文档为横竖混排（{'/'.join(orients)}）：确认横版节只用于宽表/宽图，"
                 "且之后切回了竖版。")

    sec0 = doc.sections[0]
    body_w = body_width(sec0)
    if not is_landscape(sec0) and columns_of(sec0) == 1 and body_w > 17:
        rep.warn("MARGIN", f"版心宽 {body_w:.1f}cm 偏宽，行会太长，建议左右边距 ≥3cm。")
    return body_w


def check_fonts(doc, rep, smap):
    if not resolved_eastasia("Normal", smap) and not docdefaults_eastasia(doc):
        rep.error("NORMAL_EASTASIA",
                  "Normal 样式没有 w:eastAsia 中文字体，全文中文会落到西文默认字体。")
    for sid, node in smap.items():
        if node["theme"] and node["eastAsia"]:
            rep.error("THEME_OVERRIDE",
                      f"样式「{node['name']}」同时存在 {'/'.join(node['theme'])} 和显式字体："
                      "主题属性优先，你设的字体不会生效。用 set_font() 清掉主题属性。",
                      sid)


def iter_paragraphs(doc):
    """正文段落 + 表格单元格段落，带位置标记。"""
    for i, p in enumerate(doc.paragraphs):
        yield f"body#{i}", p
    for ti, tb in enumerate(doc.tables):
        for ri, row in enumerate(tb.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    yield f"table{ti}[{ri},{ci}]", p


def is_display(p):
    """封面标题、目录标题、图注这类“故意破例”的段落，不参与正文一致性检查。"""
    name = (p.style.name or "") if p.style is not None else ""
    if any(k in name for k in ("TOC", "Title", "Subtitle", "Caption", "Heading")):
        return True
    return p.alignment is not None and int(p.alignment) in (1, 3)   # 居中 / 分散


def check_runs(doc, rep, smap):
    missing, sizes, colors = [], Counter(), Counter()
    for where, p in iter_paragraphs(doc):
        sid = p.style.style_id if p.style is not None else "Normal"
        body_text = where.startswith("body") and not is_display(p)
        inherited = resolved_eastasia(sid, smap) or docdefaults_eastasia(doc)
        for r in p.runs:
            if not r.text.strip():
                continue
            rPr = r._element.rPr
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            own = rf.get(qn("w:eastAsia")) if rf is not None else None
            theme = [a for a in THEME_ATTRS if rf is not None and rf.get(qn("w:" + a))]
            if CJK.search(r.text):
                if not own and not inherited:
                    missing.append((where, snippet(r.text)))
                elif theme and own:
                    rep.error("THEME_OVERRIDE_RUN",
                              f"run 上主题字体属性没清掉，显式字体不生效：{snippet(r.text)}", where)
            if r.font.size and body_text:
                sizes[round(r.font.size.pt, 1)] += len(r.text)
            if r.font.color is not None and r.font.color.rgb is not None:
                colors[str(r.font.color.rgb)] += 1
    if missing:
        rep.error("RUN_EASTASIA",
                  f"{len(missing)} 处中文 run 没有可继承的 eastAsia 字体，例：{missing[0][1]}",
                  missing[0][0])
    if len(sizes) > 3:
        rep.warn("SIZE_ZOO", f"正文里 run 级字号有 {len(sizes)} 种（{sorted(sizes)}），"
                             "字号应由样式统一控制，档数越多越像拼贴。")
    # 样式上定义的颜色也算数——只数 run 级颜色会漏掉“标题一个色、正文标签另一个色”
    used = {p.style.style_id for _, p in iter_paragraphs(doc) if p.style is not None}
    for sid in used:
        st = doc.styles.element.find(f'{W}style[@{W}styleId="{sid}"]')
        col = st.find(f"{W}rPr/{W}color") if st is not None else None
        if col is not None and col.get(qn("w:val")):
            colors[col.get(qn("w:val")).upper()] += 1

    accents = [c for c in colors if c.upper() not in NEUTRAL]
    if len(accents) > 1:
        rep.warn("MULTI_ACCENT",
                 f"出现 {len(accents)} 个非中性色：{accents[:5]}。全文只应锁定一个强调色。")
    if colors.get("000000", 0) > 20:
        rep.warn("PURE_BLACK", "大量正文使用纯黑 000000，长文刺眼，建议 1F2937。")
    return sizes


def check_structure(doc, rep, body_w_cm, sizes):
    empties, run_empty, heads, fake = 0, 0, [], []
    for i, p in enumerate(doc.paragraphs):
        name = (p.style.name or "") if p.style is not None else ""
        text = p.text.strip()
        xml = p._p.xml
        deliberate = any(k in xml for k in ("w:br", "w:pBdr", "w:drawing", "w:sectPr"))
        if not text and not deliberate:
            run_empty += 1
            empties = max(empties, run_empty)
            continue
        run_empty = 0
        if name.startswith("Heading"):
            try:
                heads.append((int(name.split()[-1]), text, i))
            except ValueError:
                pass
        elif (text and len(text) <= 40 and text[-1] not in END_PUNCT
              and not is_display(p)
              and p.runs and all(r.bold for r in p.runs if r.text.strip())):
            fake.append((i, snippet(text)))
        if PLACEHOLDER.search(p.text):
            rep.error("PLACEHOLDER", f"占位符残留：{snippet(p.text, 40)}", f"body#{i}")
        if "‑" in p.text or "–" in p.text:
            rep.warn("UNICODE_HYPHEN", f"含非 ASCII 连字符，渲染易出问题：{snippet(p.text)}",
                     f"body#{i}")

    if empties >= 2:
        rep.error("EMPTY_LINES", f"存在连续 {empties} 个空段落——用空行排版是返工项，"
                                 "改用 space_before / space_after。")
    elif empties == 1:
        rep.warn("EMPTY_LINES", "存在空段落，确认不是用来占位排版的。")
    for i, t in fake[:3]:
        rep.error("FAKE_HEADING",
                  f"疑似「加粗冒充标题」（Normal 样式 + 整段加粗 + 无句末标点）：{t}", f"body#{i}")
    if not heads:
        rep.warn("NO_HEADING", "全文没有 Heading 样式段落，文档没有结构、无法生成目录。")
    else:
        prev = 0
        for lvl, text, i in heads:
            if prev and lvl > prev + 1:
                rep.warn("HEADING_JUMP", f"标题层级从 H{prev} 直接跳到 H{lvl}：{snippet(text)}",
                         f"body#{i}")
            prev = lvl

    body_pt = None
    normal = doc.styles["Normal"]
    if normal.font.size:
        body_pt = normal.font.size.pt
    elif sizes:
        body_pt = sizes.most_common(1)[0][0]
    prose_chars = sum(len(p.text) for p in doc.paragraphs
                      if p.style is not None and not (p.style.name or "").startswith(
                          ("Heading", "Caption", "TOC", "Title")))
    if body_pt and prose_chars >= 200:      # 以表格为主的页面没有“行长”可言，别误报
        cm_per_char = body_pt * 2.54 / 72
        cpl = body_w_cm / cm_per_char
        # 公文（三号 16pt + 行网格）的标准就是每行 28 字，不能按商务文档的区间判
        gongwen = body_pt >= 15 and "w:docGrid" in doc.sections[0]._sectPr.xml
        lo, hi, want = (26, 32, "28（GB/T 9704）") if gongwen else (30, 46, "35-45")
        if columns_of(doc.sections[0]) > 1:      # 分栏的栏宽本来就窄，区间不同
            lo, hi, want = 20, 38, "22-35（分栏）"
        if not lo <= cpl <= hi:
            rep.warn("LINE_LENGTH",
                     f"估算每行约 {cpl:.0f} 个中文字（正文 {body_pt}pt，版心 {body_w_cm:.1f}cm），"
                     f"合适区间是 {want}，调字号或边距。")
    ls = normal.paragraph_format.line_spacing
    if ls is not None and not hasattr(ls, "pt") and ls < 1.15 and prose_chars >= 200:
        rep.warn("LINE_SPACING", f"Normal 行距 {ls} 偏挤，中文正文建议 1.3-1.5 倍。")
    return heads


def check_characters(doc, rep):
    """字符层面：乱码、直引号、夹在中文里的半角标点。

    这类问题渲染图上不显眼（一个直引号只有几个像素），但它和生成脚本的
    SyntaxError 是同一个根：中文语境里混进了 ASCII 标点。抽文字一查一个准。
    """
    bad_quote, bad_punct, garbled = [], [], []
    for where, p in iter_paragraphs(doc):
        text = p.text
        if not text.strip():
            continue
        if MOJIBAKE.search(text):
            garbled.append((where, snippet(text)))
        if CJK.search(text):
            if STRAIGHT_QUOTE.search(text):
                bad_quote.append((where, snippet(text)))
            if HALF_PUNCT.search(text):
                bad_punct.append((where, snippet(text)))
    if garbled:
        rep.error("MOJIBAKE",
                  f"{len(garbled)} 处文字含乱码或替换字符（U+FFFD / BOM / GBK 串码），"
                  f"多半是读源文件时没指定 encoding='utf-8'：{garbled[0][1]}",
                  garbled[0][0])
    if bad_quote:
        rep.warn("STRAIGHT_QUOTE",
                 f"{len(bad_quote)} 处中文里用了 ASCII 直引号 \" 或 '，应为 “ ”‘ ’："
                 f"{bad_quote[0][1]}", bad_quote[0][0])
    if bad_punct:
        rep.warn("HALF_PUNCT",
                 f"{len(bad_punct)} 处中文里夹了半角标点（, ; ! ?），中文正文应用全角："
                 f"{bad_punct[0][1]}", bad_punct[0][0])
    if bad_quote or bad_punct:
        rep.info("PUNCT_SOURCE",
                 "如果这些标点是从源文件照抄来的（Excel/台账转换类任务），"
                 "确认过就行——保真优先，别去“修”原文。")


def check_tables(doc, rep):
    for ti, tb in enumerate(doc.tables):
        where = f"table#{ti}"
        rows = len(tb.rows)
        tblPr = tb._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
        style_name = tb.style.name if tb.style is not None else ""
        if borders is None:
            if "Grid" in style_name:
                rep.error("BLACK_GRID",
                          f"表格用了「{style_name}」内置样式且没有覆盖边框——就是被禁的全黑粗网格。",
                          where)
            else:
                rep.warn("NO_BORDER", f"表格没有显式边框设置（当前样式「{style_name}」）。", where)
        else:
            for edge in ("insideH", "insideV"):
                e = borders.find(qn("w:" + edge))
                if e is None:
                    continue
                color = (e.get(qn("w:color")) or "auto").upper()
                sz = int(e.get(qn("w:sz")) or 0)
                if color in ("AUTO", "000000") and sz >= 8:
                    rep.error("BLACK_GRID", f"内框线是 {sz/8:.1f}pt 黑线，改用 0.5pt D9D9D9。", where)
                elif color in ("AUTO", "000000"):
                    rep.warn("DARK_GRID", "内框线是黑色，建议浅灰 D9D9D9。", where)

        # 表格宽度：autofit 会让 Word 忽略声明的列宽、按内容收缩——实测里表格因此只占半页
        layout = tblPr.find(qn("w:tblLayout")) if tblPr is not None else None
        tw = tblPr.find(qn("w:tblW")) if tblPr is not None else None
        fixed = layout is not None and layout.get(qn("w:type")) == "fixed"
        full = tw is not None and (tw.get(qn("w:type")) == "pct" or int(tw.get(qn("w:w")) or 0) > 0)
        if not fixed or not full:
            rep.warn("TABLE_AUTOFIT",
                     "表格是 autofit 宽度：Word 会**忽略**你设的 cell.width 按内容自行收缩，"
                     "常见结果是表格只占半页宽、右边一大片空白。"
                     "用 fit_table_to_page(t, doc=doc)（会写 tblLayout=fixed + tblW=100% + tblGrid）。",
                     where)

        if tblPr is not None and tblPr.find(qn("w:tblCellMar")) is None:
            rep.warn("CELL_MARGIN", "表格没设单元格内边距，默认上下为 0，文字贴边。", where)

        if rows:
            hdr = tb.rows[0]
            shaded = any(c._tc.find(f"{W}tcPr/{W}shd") is not None for c in hdr.cells)
            bold = any(r.bold for c in hdr.cells for p in c.paragraphs for r in p.runs)
            if not shaded and not bold:
                rep.error("BARE_HEADER", "裸表头：既没底色也没加粗。", where)
            elif not shaded:
                rep.warn("BARE_HEADER", "表头只加粗没底色，层级偏弱。", where)
            if rows > 8 and hdr._tr.find(f"{W}trPr/{W}tblHeader") is None:
                rep.warn("NO_REPEAT_HEADER", f"{rows} 行的长表格没设跨页表头重复（tblHeader）。", where)
            if rows > 7:
                shades = {c._tc.find(f"{W}tcPr/{W}shd") is not None
                          for r in tb.rows[1:] for c in r.cells}
                if shades == {False}:
                    rep.warn("NO_ZEBRA", f"{rows} 行数据表没有斑马纹，阅读时容易串行。", where)

        # 数字列对齐（同一张表合并成一条，别刷屏）
        misaligned = []
        for ci in range(len(tb.columns)):
            vals, aligns = [], []
            for ri, row in enumerate(tb.rows):
                if ri == 0:
                    continue
                cell = row.cells[ci]
                vals.append(cell.text.strip())
                aligns += [p.alignment for p in cell.paragraphs if p.text.strip()]
            numeric = [v for v in vals if v and re.fullmatch(r"[-+]?[\d,，.%％、\s]+", v)]
            if vals and len(numeric) >= max(2, len(vals) * 0.7):
                if not aligns or all(a is None or int(a) == 0 for a in aligns):
                    misaligned.append(ci + 1)
        if misaligned:
            cols = "、".join(str(c) for c in misaligned)
            rep.warn("NUM_ALIGN",
                     f"第 {cols} 列是数字列但左对齐，应右对齐"
                     f"（style_table(..., align_right_cols={tuple(c-1 for c in misaligned)})）。",
                     where)


def _text_cm(text, pt):
    """一段文字排成一行需要多宽（cm）。中文按一个字身，西文按半个。"""
    units = sum(1.0 if CJK.search(ch) else 0.5 for ch in text)
    return units * pt * 2.54 / 72


def walk_body(doc):
    """按文档顺序遍历正文块，并给出每块所属的节序号。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            yield idx, p
            if child.find(f"{W}pPr/{W}sectPr") is not None:
                idx += 1
        elif child.tag == qn("w:tbl"):
            yield idx, Table(child, doc)


def check_layout(doc, rep, sizes):
    """版式层面的检查：表格/图片是否超出版心、要不要改横版、大致多少页。

    这些问题抽文字看不出来，而在没有 LibreOffice 时也拿不到渲染页——
    所以用几何估算兜住最要命的那一类：内容比版心宽。
    """
    from docx.table import Table
    secs = doc.sections
    body_pt = doc.styles["Normal"].font.size.pt if doc.styles["Normal"].font.size else 12
    height_cm, page_est, raw_h = 0.0, 1, 0.0

    for si, block in walk_body(doc):
        sec = secs[min(si, len(secs) - 1)]
        avail = body_width(sec)
        usable_h = sec.page_height.cm - sec.top_margin.cm - sec.bottom_margin.cm

        if isinstance(block, Table):
            cells_pt = next((r.font.size.pt for row in block.rows for c in row.cells
                             for p in c.paragraphs for r in p.runs if r.font.size),
                            max(9, body_pt - 1.5))
            need = 0.0
            for col in range(len(block.columns)):
                widest = max((_text_cm(block.rows[r].cells[col].text, cells_pt)
                              for r in range(len(block.rows))), default=0)
                need += min(widest, avail / 2) + 0.4          # 0.4cm 内边距
            if need > avail * 1.15:
                fix = ("转置（指标做行、期间做列）→ 合并或删列 → 拆成两张表 → 表格字号降一档；"
                       "以上都不行且必须保留全部列时，才让这一张表单独走横版节")
                if is_landscape(sec):
                    fix = "转置、合并列或拆表——已经是横版了，再宽只能改表"
                rep.warn("TABLE_TOO_WIDE",
                         f"表格内容估算需要约 {need:.1f}cm，版心只有 {avail:.1f}cm，"
                         f"文字会大量折行。按顺序试：{fix}。"
                         "（页面方向应由文档内容性质决定，不该为一张表翻转全文）",
                         f"sec#{si}")
            widths = [c.width.cm for c in block.columns if c.width is not None]
            if len(widths) == len(block.columns):
                total_w = sum(widths)
                if total_w > avail + 0.3:
                    rep.error("TABLE_OVERFLOW",
                              f"表格固定列宽合计 {total_w:.1f}cm > 版心 {avail:.1f}cm，会超出页面。",
                              f"sec#{si}")
                elif total_w < avail * 0.75:
                    rep.warn("TABLE_UNDERFILL",
                             f"表格只有 {total_w:.1f}cm，版心 {avail:.1f}cm——"
                             f"右边会空掉 {avail - total_w:.1f}cm。fit_table_to_page() 可铺满。",
                             f"sec#{si}")
                # 列宽与内容量是否匹配：长文本列被挤窄是行高爆炸的主因。
                # 理想值要和 docx_helpers.auto_column_widths 用同一套阻尼（volume**0.65），
                # 否则会把 helper 算出的合理分配也判成失衡——按字数正比分配本来就不对。
                vols = [max(sum(_text_cm(block.rows[r].cells[c].text, 10)
                                for r in range(len(block.rows))), 0.1)
                        for c in range(len(block.columns))]
                damped = [v ** 0.65 for v in vols]
                ideal = [v / sum(damped) for v in damped]
                actual = [w / total_w for w in widths] if total_w else ideal
                worst = max(range(len(ideal)), key=lambda i: ideal[i] - actual[i])
                if ideal[worst] - actual[worst] > 0.15:
                    share = vols[worst] / sum(vols)
                    rep.warn("COLUMN_BALANCE",
                             f"第 {worst+1} 列装了全表 {share*100:.0f}% 的文字，"
                             f"只分到 {actual[worst]*100:.0f}% 的宽度（合理值约 "
                             f"{ideal[worst]*100:.0f}%）——这一列会疯狂折行、把行高撑得很高。"
                             "用 fit_table_to_page() 按内容量分配列宽。",
                             f"sec#{si}")
            # 行高要按“单元格里折了几行”算，不能按行数×一行——长文本表格差几倍
            col_w = [c.width.cm if c.width is not None else avail / max(1, len(block.columns))
                     for c in block.columns]
            line_h = cells_pt * 2.54 / 72 * LINE_H * CALIBRATION
            for row in block.rows:
                lines = 1
                for ci, cell in enumerate(row.cells):
                    w = max(col_w[ci] if ci < len(col_w) else avail, 0.5) - 0.4
                    for para in cell.paragraphs:
                        lines = max(lines, max(1, round(_text_cm(para.text, cells_pt) / w + WRAP_BIAS)))
                height_cm += lines * line_h + 0.05
                raw_h += lines * line_h + 0.05
            height_cm += 0.6
            raw_h += 0.6
        else:
            xml = block._p.xml
            for m in re.finditer(r'<wp:extent cx="(\d+)" cy="(\d+)"', xml):
                w_cm, h_cm = int(m.group(1)) / 360000, int(m.group(2)) / 360000
                if w_cm > avail + 0.2:
                    rep.error("IMAGE_OVERFLOW",
                              f"图片宽 {w_cm:.1f}cm 超出版心 {avail:.1f}cm，会被裁掉或撑破版面。",
                              f"sec#{si}")
                height_cm += h_cm + 0.8
                raw_h += h_cm + 0.8
            if "w:br" in xml and 'w:type="page"' in xml:
                page_est += 1
                height_cm = 0
                continue
            if block.text.strip():
                pt = next((r.font.size.pt for r in block.runs if r.font.size), body_pt)
                lines = max(1, round(_text_cm(block.text, pt) / avail + WRAP_BIAS))
                pf = block.paragraph_format
                extra = ((pf.space_before.pt if pf.space_before else 0)
                         + (pf.space_after.pt if pf.space_after else 8)) * 2.54 / 72
                height_cm += lines * pt * 2.54 / 72 * 1.45 + extra
                raw_h += lines * pt * 2.54 / 72 * 1.45 + extra
        if height_cm > usable_h * 1.05:      # 5% 容差：估算误差就在这个量级，别在边界上抖
            page_est += int(height_cm // usable_h)
            height_cm = height_cm % usable_h

    sec0 = secs[0]
    usable_h0 = sec0.page_height.cm - sec0.top_margin.cm - sec0.bottom_margin.cm
    pages_f = raw_h / usable_h0 if usable_h0 else 0
    if page_est == 1:
        fill = min(height_cm, raw_h) / usable_h0
        if fill < 0.6:
            rep.info("PAGE_FILL",
                     f"内容大约只占页面高度的 {fill*100:.0f}%。"
                     "要求“铺满一页”时：先确认表格已铺满版心，再加大字号/行距/内边距。"
                     "别用空行去填，也别为了填版面去增删内容。")
    autofit = any(t._tbl.tblPr is not None
                  and (t._tbl.tblPr.find(qn("w:tblLayout")) is None
                       or t._tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")) != "fixed")
                  for t in doc.tables)
    rep.info("PAGE_ESTIMATE",
             f"粗估约 {page_est} 页，内容高度 ≈ {pages_f:.2f} 个版心（几何估算，误差约 ±15%）。"
             + ("表格是 autofit，实际列宽由 Word 现算，这个估算偏差会很大——"
                "先用 fit_table_to_page() 固定列宽再看。" if autofit else
                "页数超出预期时，按“边距 → 列宽分配 → 字号（下限 8pt）→ 行距”的顺序调版面；"
                "**不要改写或删减原文内容**，实在放不下就如实告诉用户。"))
    return page_est


def check_fields(doc, rep, heads):
    xml = doc.element.xml
    footers = " ".join(
        p._p.xml for s in doc.sections
        for f in (s.footer, s.first_page_footer, s.even_page_footer) if f is not None
        for p in f.paragraphs
    )
    if "PAGE" not in footers:
        rep.warn("NO_PAGE_NUMBER", "页脚没有页码域，正式交付必须有页码。")
    if len(heads) >= 8 and "TOC" not in xml and "目录" not in xml:
        rep.warn("NO_TOC", f"有 {len(heads)} 个标题但没有目录，>5 节的长文应加目录。")
    if "TOC \\o" in xml or 'TOC \\o' in xml:
        rep.info("TOC_FIELD",
                 "目录是 Word 域：LibreOffice 转出的 PDF 里不会有真页码，"
                 "交付时要告诉用户在 Word 中打开会自动更新。")
    # 以表格为主、正文极少 → 多半是 Excel/台账转换，内容保真比排版更要紧
    prose = sum(len(p.text) for p in doc.paragraphs)
    if doc.tables and prose < 300:
        rep.info("CONTENT_FIDELITY",
                 "这是一份以表格为主的文档（像是从 Excel/台账转来的）："
                 "交付前请逐单元格与源文件比对，确认没有改写、缩写或漏行——"
                 "转换类任务不允许为了排版去动内容。")

    pics = xml.count("<w:drawing>")
    if pics:
        caps = sum(1 for _, p in iter_paragraphs(doc)
                   if p.style is not None and (p.style.name or "").startswith("Caption"))
        if caps < pics:
            rep.warn("NO_CAPTION", f"{pics} 张图里只有 {caps} 条图注，图注应含图名与数据来源。")
    return pics


def main():
    ap = argparse.ArgumentParser(description="DOCX 交付前自检")
    ap.add_argument("path")
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--letter", action="store_true", help="允许 US Letter 页面")
    ap.add_argument("--strict", action="store_true", help="WARN 也算失败")
    args = ap.parse_args()

    doc = Document(args.path)
    rep = Report()
    smap = style_font_map(doc)

    body_w = check_page(doc, rep, args.letter)
    check_fonts(doc, rep, smap)
    sizes = check_runs(doc, rep, smap)
    heads = check_structure(doc, rep, body_w, sizes)
    check_characters(doc, rep)
    check_tables(doc, rep)
    check_layout(doc, rep, sizes)
    pics = check_fields(doc, rep, heads)

    errors = [i for i in rep.items if i["level"] == "ERROR"]
    warns = [i for i in rep.items if i["level"] == "WARN"]

    if args.json:
        print(json.dumps({"path": args.path, "errors": len(errors),
                          "warnings": len(warns), "items": rep.items},
                         ensure_ascii=False, indent=2))
    else:
        icon = {"ERROR": "×", "WARN": "!", "INFO": "·"}
        for lvl in ("ERROR", "WARN", "INFO"):
            for it in [i for i in rep.items if i["level"] == lvl]:
                loc = f" [{it['where']}]" if it["where"] else ""
                print(f"{icon[lvl]} {lvl:<5} {it['code']}{loc}: {it['msg']}")
        print(f"\n{len(errors)} error(s), {len(warns)} warning(s) — "
              f"{len(doc.paragraphs)} 段, {len(doc.tables)} 表, {pics} 图")
        if not rep.items:
            print("干净。仍需渲染成 PNG 逐页看分页与溢出。")

    sys.exit(1 if errors or (args.strict and warns) else 0)


if __name__ == "__main__":
    main()
